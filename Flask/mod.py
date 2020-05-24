import datetime
import os
import pyowm

from Flask import app
from table_class import user as clUser,worker as worker_py,kvalif_up as kvalif_up_py, retraining as retraining_py, \
    vac_certification as vac_certification_py, vacation as vacation_py, \
    assignment_and_relocation as assignment_and_relocation_py, staff_list as staff_list_py
import docx
from docxtpl import DocxTemplate
workerlists_dir = 'Flask/static/files/workerlists/'
personalList_dir = 'Flask/static/files/personalLists/'
staffList_dir = 'Flask/static/files/staffLists/'


ALLOWED_EXTENSIONS = {'png'}

def show_alert(alert_type,header, body):
    """
    def x():
        alert_html = request.args.get('alert_html')
        if alert_html is None:
            alert_html = ''

    return render_template('x.html',alert_html=alert_html)


     alert_html = Flask.mod.show_alert('success', 'Отлично! ', 'данные сохранены')
    return flask.redirect(flask.url_for('x', alert_html=alert_html))
    """
    #success
    #info
    #warning
    #danger
    html = '<div class="showback"><div class="alert alert-'+alert_type+'">' \
           '<b>'+header+'</b> ' \
           ''+body+'</div></div>'
    return html



def genWorkerList(worker_id, session):
    wor = session.query(worker_py.Worker).filter_by(id=worker_id).first()
    vac_cert = session.query(vac_certification_py.Vac_certification).filter_by(worker_id=worker_id).all()
    vacat = session.query(vacation_py.Vacation).filter_by(worker_id=worker_id).all()
    ret = session.query(retraining_py.Retraining).filter_by(worker_id=worker_id).all()
    kvalif = session.query(kvalif_up_py.Kvalif_up).filter_by(worker_id=worker_id).all()
    as_reloc = session.query(assignment_and_relocation_py.Assignment_and_relocation).filter_by(worker_id=worker_id).all()
    doc = DocxTemplate("Flask/static/templates/workerlists.docx")
    context = worker_py.worker_serializer(wor)
    if wor.gender_id == 0:
        context['gender_name'] = "Женский"
    else:
        context['gender_name'] = "Мужской"
    if vac_cert != None:
        context['vac_cert'] = vac_cert
    if vacat != None:
        context['vacat'] = vacat
    if ret != None:
        context['ret'] = ret
    if kvalif != None:
        context['kvalif'] = kvalif
    if as_reloc != None:
        context['as_reloc'] = as_reloc

    doc.render(context)
    name = str(wor.id)+" - "+str(datetime.datetime.now().strftime("%d-%m-%Y %H-%M-%S")) +".docx"
    path = workerlists_dir + name
    doc.save(path)



    """
    docOnTable = docx.Document(path)
    table = doc.add_table(rows=3, cols=3)
    #table = docOnTable.table[0]
    for row in range(3):
        for col in range(3):
            # получаем ячейку таблицы
            cell = table.cell(row, col)
            # записываем в ячейку данные
            cell.text = str(row + 1) + str(col + 1)

    docOnTable.save(path)
"""

    session.close()
    return name

def genPersonalList(worker_id, session):
    wor = session.query(worker_py.Worker).filter_by(id=worker_id).first()

    doc = DocxTemplate("Flask/static/templates/personalList.docx")
    context = worker_py.worker_serializer(wor)
    if wor.gender_id == 0:
        context['gender_name'] = "Женский"
    else:
        context['gender_name'] = "Мужской"


    doc.render(context)
    name = str(wor.id)+" - "+str(datetime.datetime.now().strftime("%d-%m-%Y %H-%M-%S")) +".docx"
    path = personalList_dir + name
    doc.save(path)
    session.close()
    return name

def genStaffList(session):
    stf = session.query(staff_list_py.Staff_list).all()
    doc = DocxTemplate("Flask/static/templates/staffList.docx")
    context = {}
    context['staff_list'] = stf
    import locale
    locale.setlocale(locale.LC_ALL, "")

    now = datetime.datetime.now()

    context['day'] = datetime.datetime.today().day
    context['month'] = now.strftime("%B")
    context['year'] =  datetime.datetime.today().year

    sum_kol_vo = session.execute('select sum(kol_vo) as sum from staff_list').first()
    context['sum_kol_vo'] = sum_kol_vo[0]
    fond = session.execute('select sum(fond) as sum from staff_list').first()
    context['sum_fond'] = fond[0]

    doc.render(context)
    name = "Штатное расписание - "+str(datetime.datetime.now().strftime("%d-%m-%Y %H-%M-%S")) +".docx"
    path = staffList_dir + name
    doc.save(path)
    session.close()
    return name

def getWeather():
    apiKey = "12bf12b5303397b18dce8a41f3a72b1d"
    import pyowm

    owm = pyowm.OWM(apiKey, language = "RU")  # You MUST provide a valid API key

    # Have a pro subscription? Then use:
    # owm = pyowm.OWM(API_key='your-API-key', subscription_type='pro')

    # Search for current weather in London (Great Britain)
    observation = owm.weather_at_place('Nur-Sultan,KZ')
    w = observation.get_weather()
    print(w)  # <Weather - reference time=2013-12-18 09:20,
    # status=Clouds>

    # Weather details
    w.get_wind()  # {'speed': 4.6, 'deg': 330}
    w.get_humidity()  # 87
    w.get_temperature('celsius')  # {'temp_max': 10.5, 'temp': 9.7, 'temp_min': 9.0}

    # Search current weather observations in the surroundings of
    # lat=22.57W, lon=43.12S (Rio de Janeiro, BR)
    observation_list = owm.weather_around_coords(-22.57, -43.12)
    return  w

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1] in ALLOWED_EXTENSIONS



def isint(s):
    try:
        int(s)
        return True
    except ValueError:
        return False

def test():





    doc = docx.Document('generated.docx')

    # добавляем таблицу 3x3
    table = doc.add_table(rows=3, cols=3)
    # применяем стиль для таблицы

    # заполняем таблицу даннымиц
    for row in range(3):
        for col in range(3):
            # получаем ячейку таблицы
            cell = table.cell(row, col)
            # записываем в ячейку данные
            cell.text = str(row + 1) + str(col + 1)

    doc.save('generated.docx')

